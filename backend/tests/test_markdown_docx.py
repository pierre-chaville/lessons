from io import BytesIO
import re

from docx import Document
from docx.oxml.ns import qn

from services.markdown_docx import docx_bytes_to_markdown, markdown_to_docx_bytes

_HEBREW_CHAR_RE = re.compile(r"[\u0590-\u05FF]")
_BIDI_CONTROL_RE = re.compile(r"[\u200E\u200F\u202A-\u202E\u2066-\u2069]")


def _has_page_break(paragraph) -> bool:
    for run in paragraph.runs:
        for br in run._r.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def _run_has_rtl(run) -> bool:
    r_pr = run._r.find(qn("w:rPr"))
    if r_pr is None:
        return False
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        return False
    return rtl.get(qn("w:val")) == "1"


def _paragraph_has_bidi(paragraph) -> bool:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return False
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        return False
    return bidi.get(qn("w:val")) == "1"


def _strip_bidi_controls(text: str) -> str:
    return _BIDI_CONTROL_RE.sub("", text or "")


def test_horizontal_rule_creates_page_break():
    markdown = "# Title\n\nParagraph before break.\n\n---\n\n## After break"
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    assert any(_has_page_break(p) for p in doc.paragraphs)
    assert any("After break" in p.text for p in doc.paragraphs)


def test_nested_bullets_have_increasing_indentation():
    markdown = (
        "- **Chapitre 1**\n"
        "- 2025-12-20 - Midrash...\n"
        "  - Nested level 1\n"
        "    - Nested level 2\n"
    )
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    bullet_paragraphs = [p for p in doc.paragraphs if p.text.strip().startswith("• ")]
    assert len(bullet_paragraphs) >= 4

    indents = [p.paragraph_format.left_indent.pt if p.paragraph_format.left_indent else 0 for p in bullet_paragraphs[:4]]
    assert indents[0] == indents[1]
    assert indents[2] > indents[1]
    assert indents[3] > indents[2]


def test_ordered_list_keeps_explicit_marker_and_indentation():
    markdown = "1. First\n  2. Nested second\n"
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    ordered = [p for p in doc.paragraphs if p.text.strip()]
    assert ordered[0].text.startswith("1. ")
    assert ordered[1].text.startswith("2. ")
    assert ordered[1].paragraph_format.left_indent.pt > ordered[0].paragraph_format.left_indent.pt


def test_ordered_list_repeated_one_markers_are_numbered_in_docx():
    markdown = "1. First\n1. Second\n1. Third\n"
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    ordered = [p.text for p in doc.paragraphs if p.text.strip()]
    assert ordered == ["1. First", "2. Second", "3. Third"]


def test_ordered_list_repeated_markers_keep_start_number():
    markdown = "4. Fourth\n4. Fifth\n"
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    ordered = [p.text for p in doc.paragraphs if p.text.strip()]
    assert ordered == ["4. Fourth", "5. Fifth"]


def test_mixed_nested_lists_keep_order_and_indentation():
    markdown = (
        "- Chapter root\n"
        "  1. Numbered child\n"
        "    - Bullet grandchild\n"
        "  2. Numbered sibling\n"
        "- Next root\n"
    )
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    lines = [p for p in doc.paragraphs if p.text.strip()]
    assert lines[0].text.startswith("• ")
    assert lines[1].text.startswith("1. ")
    assert lines[2].text.startswith("• ")
    assert lines[3].text.startswith("2. ")
    assert lines[4].text.startswith("• ")

    indents = [p.paragraph_format.left_indent.pt if p.paragraph_format.left_indent else 0 for p in lines[:5]]
    assert indents[1] > indents[0]
    assert indents[2] > indents[1]
    assert indents[3] == indents[1]
    assert indents[4] == indents[0]


def test_docx_to_markdown_round_trip_supported_structure():
    markdown = (
        "# Report title\n\n"
        "Intro with **bold** and *italic* text.\n\n"
        "- Bullet root\n"
        "  - Nested bullet\n"
        "  1. Nested ordered\n\n"
        "> A quoted thought\n\n"
        "---\n\n"
        "## After break\n"
    )
    docx_bytes = markdown_to_docx_bytes(markdown)

    converted = docx_bytes_to_markdown(docx_bytes)

    assert "# Report title" in converted
    assert "Intro with **bold** and *italic* text." in converted
    assert "- Bullet root" in converted
    assert "  - Nested bullet" in converted
    assert "  1. Nested ordered" in converted
    assert "> A quoted thought" in converted
    assert "---" in converted
    assert "## After break" in converted


def test_docx_to_markdown_keeps_list_indentation_from_exporter():
    markdown = "- Root\n  - Child\n    - Grandchild\n"
    docx_bytes = markdown_to_docx_bytes(markdown)

    converted = docx_bytes_to_markdown(docx_bytes)

    assert "- Root" in converted
    assert "  - Child" in converted
    assert "    - Grandchild" in converted


def test_docx_to_markdown_imports_native_numbered_list_style():
    document = Document()
    document.add_paragraph("First", style="List Number")
    document.add_paragraph("Second", style="List Number")

    output = BytesIO()
    document.save(output)

    converted = docx_bytes_to_markdown(output.getvalue())

    assert converted == "1. First\n2. Second"


def test_docx_round_trip_preserves_section_markers():
    markdown = (
        "# Meta\n\n"
        "<!-- MARKER:section-start -->\n\n"
        "Summary body\n\n"
        "<!-- MARKER:section-end -->\n"
    )
    docx_bytes = markdown_to_docx_bytes(markdown)

    converted = docx_bytes_to_markdown(docx_bytes)

    assert "<!-- MARKER:section-start -->" in converted
    assert "<!-- MARKER:section-end -->" in converted


def test_docx_round_trip_preserves_italic_heading_segments():
    markdown = "## *Matan Torathenou* et *Qabbalat ha-Torah*"

    docx_bytes = markdown_to_docx_bytes(markdown)
    converted = docx_bytes_to_markdown(docx_bytes)

    assert "## *Matan Torathenou* et *Qabbalat ha-Torah*" in converted


def test_docx_heading_import_ignores_heading_bold_noise_in_italic_runs():
    document = Document()
    heading = document.add_heading(level=2)

    run_one = heading.add_run("Matan Torathenou")
    run_one.italic = True
    heading.add_run(" et ")
    run_two = heading.add_run("Qabbalat")
    run_two.italic = True
    run_three = heading.add_run(" ha-Torah")
    run_three.italic = True
    run_three.bold = True

    output = BytesIO()
    document.save(output)

    converted = docx_bytes_to_markdown(output.getvalue())

    assert converted == "## *Matan Torathenou* et *Qabbalat ha-Torah*"


def test_docx_import_escapes_literal_markdown_control_chars():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run(r"Price *value* and key_name ")
    italic = paragraph.add_run(r"*wrapped* _token_")
    italic.italic = True

    output = BytesIO()
    document.save(output)

    converted = docx_bytes_to_markdown(output.getvalue())

    assert converted == r"Price \*value\* and key\_name *\*wrapped\* \_token\_*"


def test_docx_import_moves_trailing_space_outside_closing_emphasis_markers():
    document = Document()
    heading = document.add_heading(level=2)
    lead = heading.add_run("De Mara au Sinaï : ce qui change avec ")
    lead.bold = True
    tail = heading.add_run("Matan Torah ")
    tail.bold = True
    tail.italic = True

    output = BytesIO()
    document.save(output)

    converted = docx_bytes_to_markdown(output.getvalue())

    assert "Matan Torah *" not in converted
    assert converted.endswith("*Matan Torah*")


def test_hebrew_runs_are_explicitly_marked_rtl():
    markdown = "Texte mixte: שלוש רגלים et encore."
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    paragraph = next(p for p in doc.paragraphs if p.text.strip())
    hebrew_runs = [run for run in paragraph.runs if _HEBREW_CHAR_RE.search(run.text or "")]

    assert hebrew_runs
    assert all(_run_has_rtl(run) for run in hebrew_runs)


def test_predominantly_hebrew_paragraph_sets_bidi():
    markdown = "שלוש רגלים ועוד מילים בעברית, avec un peu."
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    paragraph = next(p for p in doc.paragraphs if p.text.strip())
    assert _paragraph_has_bidi(paragraph)


def test_mixed_french_hebrew_marks_only_hebrew_runs_rtl():
    markdown = (
        "Dans la tefila, les שלוש רגלים (shalosh regalim, "
        "« les trois fêtes de pèlerinage ») sont appelées ..."
    )
    docx_bytes = markdown_to_docx_bytes(markdown)
    doc = Document(BytesIO(docx_bytes))

    paragraph = next(p for p in doc.paragraphs if p.text.strip())
    assert _strip_bidi_controls(paragraph.text) == markdown

    hebrew_runs = [run for run in paragraph.runs if _HEBREW_CHAR_RE.search(run.text or "")]
    latin_runs = [run for run in paragraph.runs if (run.text or "").strip() and not _HEBREW_CHAR_RE.search(run.text or "")]

    assert hebrew_runs
    assert latin_runs
    assert all(_run_has_rtl(run) for run in hebrew_runs)
    assert all(not _run_has_rtl(run) for run in latin_runs)
