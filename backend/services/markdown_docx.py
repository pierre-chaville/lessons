"""Utilities to convert simple markdown documents into DOCX."""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

_HEADING_RE = re.compile(r"^\s{0,3}(#{1,6})\s+(.+)$")
_HORIZONTAL_RULE_RE = re.compile(r"^\s{0,3}((\*\s*){3,}|(-\s*){3,}|(_\s*){3,})\s*$")
_UNORDERED_LIST_RE = re.compile(r"^(?P<indent>\s*)[-*+]\s+(?P<content>.+)$")
_ORDERED_LIST_RE = re.compile(r"^(?P<indent>\s*)(?P<number>\d+)[.)]\s+(?P<content>.+)$")
_BLOCKQUOTE_RE = re.compile(r"^\s{0,3}(?P<markers>(>\s*)+)(?P<content>.*)$")
_HTML_COMMENT_LINE_RE = re.compile(r"^\s*<!--.*-->\s*$")
_INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_)")
_DOCX_ORDERED_MARKER_RE = re.compile(r"^(?P<number>\d+)\.\s+(?P<content>.+)$")
_SECTION_START_MARKER = "<!-- MARKER:section-start -->"
_SECTION_END_MARKER = "<!-- MARKER:section-end -->"
_DOCX_SECTION_START_SENTINEL = "[[LESSONS_SECTION_START]]"
_DOCX_SECTION_END_SENTINEL = "[[LESSONS_SECTION_END]]"
_HEBREW_CHAR_RE = re.compile(r"[\u0590-\u05FF]")
_HEBREW_WORD_RE = r"[\u0590-\u05FF\u05BE\u05F3\u05F4]+"
_HEBREW_SPAN_RE = re.compile(rf"{_HEBREW_WORD_RE}(?:\s+{_HEBREW_WORD_RE})*")
_LATIN_CHAR_RE = re.compile(r"[A-Za-z]")
_BIDI_CONTROL_RE = re.compile(r"[\u200E\u200F\u202A-\u202E\u2066-\u2069]")


def _escape_markdown_inline_text(text: str) -> str:
    """Escape markdown control characters that should remain literal."""
    escaped = str(text or "")
    escaped = escaped.replace("\\", "\\\\")
    escaped = escaped.replace("*", "\\*")
    escaped = escaped.replace("_", "\\_")
    return escaped


def _is_predominantly_hebrew(text: str) -> bool:
    content = str(text or "")
    hebrew_count = len(_HEBREW_CHAR_RE.findall(content))
    latin_count = len(_LATIN_CHAR_RE.findall(content))
    return hebrew_count > 0 and hebrew_count >= latin_count


def _set_run_rtl(run) -> None:
    # Mixed French/Hebrew text in Word is stable when Hebrew spans are marked
    # as RTL + complex script + Hebrew bidi language at the run level.
    # We avoid Unicode control characters in content because Word can surface
    # them visually in some display/debug modes.
    r_pr = run._r.get_or_add_rPr()
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")

    # Mark the run as complex-script Hebrew to improve Word's line-wrap behavior
    # for mixed-direction paragraphs without injecting visible control symbols.
    cs = r_pr.find(qn("w:cs"))
    if cs is None:
        cs = OxmlElement("w:cs")
        r_pr.append(cs)
    cs.set(qn("w:val"), "1")

    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:bidi"), "he-IL")


def _set_paragraph_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _add_run_with_direction(paragraph, text: str, *, bold: bool = False, italic: bool = False, rtl: bool = False):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if rtl:
        _set_run_rtl(run)
    return run


def _append_text_with_direction(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    cursor = 0
    for match in _HEBREW_SPAN_RE.finditer(text):
        start, end = match.span()
        if start > cursor:
            _add_run_with_direction(paragraph, text[cursor:start], bold=bold, italic=italic)
        _add_run_with_direction(paragraph, text[start:end], bold=bold, italic=italic, rtl=True)
        cursor = end

    if cursor < len(text):
        _add_run_with_direction(paragraph, text[cursor:], bold=bold, italic=italic)


def _finalize_paragraph_direction(paragraph) -> None:
    if _is_predominantly_hebrew(paragraph.text):
        _set_paragraph_bidi(paragraph)


def _append_inline_runs(paragraph, text: str) -> None:
    """Append text to a paragraph while handling bold/italic markdown markers."""
    cursor = 0
    for match in _INLINE_TOKEN_RE.finditer(text):
        start, end = match.span()
        if start > cursor:
            _append_text_with_direction(paragraph, text[cursor:start])

        token = match.group(0)
        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            _append_text_with_direction(paragraph, token[2:-2], bold=True)
        elif token.startswith("__") and token.endswith("__") and len(token) >= 4:
            _append_text_with_direction(paragraph, token[2:-2], bold=True)
        elif token.startswith("*") and token.endswith("*") and len(token) >= 2:
            _append_text_with_direction(paragraph, token[1:-1], italic=True)
        elif token.startswith("_") and token.endswith("_") and len(token) >= 2:
            _append_text_with_direction(paragraph, token[1:-1], italic=True)
        else:
            _append_text_with_direction(paragraph, token)
        cursor = end

    if cursor < len(text):
        _append_text_with_direction(paragraph, text[cursor:])


def _list_indent_level(indent: str) -> int:
    # Tabs are treated as 4 spaces; every 2 spaces increases nesting level.
    expanded = indent.replace("\t", "    ")
    return max(0, len(expanded) // 2)


def _add_list_paragraph(
    doc: Document,
    *,
    indent_level: int,
    marker: str,
    content: str,
) -> None:
    """
    Add a list paragraph with explicit indentation.

    We intentionally render list markers ourselves rather than relying on
    Word list styles, because style availability/behavior can vary by template.
    """
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    base_indent = 18
    hanging_indent = 12
    paragraph_format.left_indent = Pt(base_indent * (indent_level + 1))
    paragraph_format.first_line_indent = Pt(-hanging_indent)
    paragraph.add_run(f"{marker} ")
    _append_inline_runs(paragraph, content.strip())
    _finalize_paragraph_direction(paragraph)


def _parse_list_item(line: str) -> tuple[int, str, str] | None:
    unordered = _UNORDERED_LIST_RE.match(line)
    if unordered:
        indent_level = _list_indent_level(unordered.group("indent"))
        return indent_level, "•", unordered.group("content")

    ordered = _ORDERED_LIST_RE.match(line)
    if ordered:
        indent_level = _list_indent_level(ordered.group("indent"))
        marker = f"{ordered.group('number')}."
        return indent_level, marker, ordered.group("content")

    return None


def _parse_blockquote_line(line: str) -> tuple[int, str] | None:
    match = _BLOCKQUOTE_RE.match(line)
    if not match:
        return None
    depth = match.group("markers").count(">")
    content = match.group("content").strip()
    return depth, content


def _is_ignored_comment_line(line: str) -> bool:
    return bool(_HTML_COMMENT_LINE_RE.match(line))


def _is_section_marker_line(line: str) -> bool:
    return str(line or "").strip() in {_SECTION_START_MARKER, _SECTION_END_MARKER}


def _section_sentinel_for_line(line: str) -> str:
    stripped = str(line or "").strip()
    if stripped == _SECTION_START_MARKER:
        return _DOCX_SECTION_START_SENTINEL
    if stripped == _SECTION_END_MARKER:
        return _DOCX_SECTION_END_SENTINEL
    return ""


def _add_blockquote_paragraph(doc: Document, *, depth: int, content: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(20 * max(depth, 1))
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.space_before = Pt(3)
    paragraph_format.space_after = Pt(3)
    _append_inline_runs(paragraph, content)
    _finalize_paragraph_direction(paragraph)


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """
    Convert a simple markdown string to DOCX bytes.

    Supported markdown features:
    - Headings (# to ######)
    - Horizontal rules (---, ***, ___) as page breaks
    - Bulleted lists (-, *, +)
    - Numbered lists (1. / 1))
    - Block quotes (> ...)
    - Paragraphs
    - Inline bold/italic (**text**, __text__, *text*, _text_)
    - HTML comment markers on standalone lines are ignored
    """
    doc = Document()
    lines = str(markdown_text or "").splitlines()
    idx = 0

    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if not stripped:
            idx += 1
            continue

        if _is_section_marker_line(raw):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(_section_sentinel_for_line(raw))
            # Keep section markers in DOCX for reliable round-trip, without visible noise.
            run.font.hidden = True
            idx += 1
            continue

        if _is_ignored_comment_line(raw):
            idx += 1
            continue

        if _HORIZONTAL_RULE_RE.match(raw):
            page_break_paragraph = doc.add_paragraph()
            page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)
            idx += 1
            continue

        heading_match = _HEADING_RE.match(raw)
        if heading_match:
            level = min(len(heading_match.group(1)), 6)
            heading_text = heading_match.group(2).strip()
            paragraph = doc.add_heading(level=level)
            _append_inline_runs(paragraph, heading_text)
            _finalize_paragraph_direction(paragraph)
            idx += 1
            continue

        if _parse_blockquote_line(raw):
            quote_lines: list[str] = []
            quote_depth = 1
            while idx < len(lines):
                line = lines[idx]
                if _is_ignored_comment_line(line):
                    idx += 1
                    continue

                line_stripped = line.strip()
                if not line_stripped:
                    idx += 1
                    break

                parsed_quote = _parse_blockquote_line(line)
                if not parsed_quote:
                    break

                line_depth, quote_content = parsed_quote
                quote_depth = max(quote_depth, line_depth)
                if quote_content:
                    quote_lines.append(quote_content)
                idx += 1

            quote_text = " ".join(quote_lines).strip()
            if quote_text:
                _add_blockquote_paragraph(doc, depth=quote_depth, content=quote_text)
            continue

        if _parse_list_item(raw):
            while idx < len(lines):
                line = lines[idx]
                if _is_ignored_comment_line(line):
                    idx += 1
                    continue

                line_stripped = line.strip()
                if not line_stripped:
                    idx += 1
                    break

                parsed = _parse_list_item(line)
                if not parsed:
                    break
                indent_level, marker, content = parsed
                _add_list_paragraph(
                    doc,
                    indent_level=indent_level,
                    marker=marker,
                    content=content,
                )
                idx += 1
            continue

        paragraph_lines: list[str] = []
        while idx < len(lines):
            line = lines[idx]
            if _is_ignored_comment_line(line):
                idx += 1
                continue

            line_stripped = line.strip()
            if not line_stripped:
                idx += 1
                break
            if (
                _HEADING_RE.match(line)
                or _UNORDERED_LIST_RE.match(line)
                or _ORDERED_LIST_RE.match(line)
                or _HORIZONTAL_RULE_RE.match(line)
                or _BLOCKQUOTE_RE.match(line)
            ):
                break
            paragraph_lines.append(line_stripped)
            idx += 1

        paragraph_text = " ".join(paragraph_lines).strip()
        if paragraph_text:
            paragraph = doc.add_paragraph()
            _append_inline_runs(paragraph, paragraph_text)
            _finalize_paragraph_direction(paragraph)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _is_page_break_paragraph(paragraph) -> bool:
    for run in paragraph.runs:
        for br in run._r.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def _inline_runs_to_markdown(paragraph, *, suppress_bold: bool = False) -> str:
    parts: list[str] = []
    active_bold = False
    active_italic = False

    def _open_markers(want_bold: bool, want_italic: bool) -> str:
        return ("**" if want_bold else "") + ("*" if want_italic else "")

    def _close_markers(close_bold: bool, close_italic: bool) -> str:
        # Close in reverse order to keep markdown nesting balanced.
        return ("*" if close_italic else "") + ("**" if close_bold else "")

    def _append_closing_markers_preserving_spacing(markers: str) -> None:
        """
        Keep trailing whitespace outside closing emphasis markers.

        External DOCX often stores words as styled runs ending with a space.
        If we emit the closing marker after that space, markdown emphasis breaks
        (e.g. `*word *`). We move trailing spaces after the marker instead
        (e.g. `*word* `).
        """
        if not markers:
            return
        trailing_ws = ""
        if parts:
            match = re.search(r"\s+$", parts[-1])
            if match:
                trailing_ws = match.group(0)
                parts[-1] = parts[-1][: -len(trailing_ws)]
        parts.append(markers)
        if trailing_ws:
            parts.append(trailing_ws)

    for run in paragraph.runs:
        text = (run.text or "").replace("\r", "")
        text = _BIDI_CONTROL_RE.sub("", text)
        if not text:
            continue
        text = _escape_markdown_inline_text(text)
        want_bold = bool(run.bold) and not suppress_bold
        want_italic = bool(run.italic)

        # Close styles that are no longer active.
        _append_closing_markers_preserving_spacing(
            _close_markers(
                close_bold=active_bold and not want_bold,
                close_italic=active_italic and not want_italic,
            )
        )
        # Open newly active styles.
        parts.append(
            _open_markers(
                want_bold=want_bold and not active_bold,
                want_italic=want_italic and not active_italic,
            )
        )
        parts.append(text)
        active_bold = want_bold
        active_italic = want_italic

    if active_bold or active_italic:
        _append_closing_markers_preserving_spacing(
            _close_markers(close_bold=active_bold, close_italic=active_italic)
        )

    return "".join(parts).strip()


def _docx_indent_level(paragraph, *, base_indent: int) -> int:
    left_indent = paragraph.paragraph_format.left_indent
    if not left_indent:
        return 0
    return max(0, round(left_indent.pt / base_indent) - 1)


def _is_docx_blockquote(paragraph) -> bool:
    left_indent = paragraph.paragraph_format.left_indent
    if not left_indent or left_indent.pt < 20:
        return False
    first_line_indent = paragraph.paragraph_format.first_line_indent
    if first_line_indent and abs(first_line_indent.pt) > 0.01:
        return False
    return True


def docx_bytes_to_markdown(docx_bytes: bytes) -> str:
    """
    Convert DOCX bytes into simple markdown.

    Supported DOCX features mirror markdown_to_docx_bytes:
    - Headings
    - Page breaks converted to horizontal rules (---)
    - List paragraphs emitted by markdown_to_docx_bytes
    - Block quote paragraphs emitted by markdown_to_docx_bytes
    - Paragraphs
    - Inline bold/italic
    """
    doc = Document(BytesIO(docx_bytes))
    blocks: list[tuple[str, str]] = []

    for paragraph in doc.paragraphs:
        if _is_page_break_paragraph(paragraph):
            blocks.append(("rule", "---"))
            continue

        raw_text = (paragraph.text or "").strip()

        style_name = ""
        if paragraph.style is not None:
            style_name = str(paragraph.style.name or "").strip().lower()

        heading_match = re.match(r"heading\s+([1-6])$", style_name)
        inline_text = _inline_runs_to_markdown(paragraph, suppress_bold=bool(heading_match))
        if not raw_text and not inline_text:
            continue

        if raw_text in {_DOCX_SECTION_START_SENTINEL, _DOCX_SECTION_END_SENTINEL}:
            marker = _SECTION_START_MARKER if raw_text == _DOCX_SECTION_START_SENTINEL else _SECTION_END_MARKER
            blocks.append(("marker", marker))
            continue

        if heading_match:
            level = int(heading_match.group(1))
            blocks.append(("heading", f"{'#' * level} {inline_text or raw_text}"))
            continue

        list_level = _docx_indent_level(paragraph, base_indent=18)
        if raw_text.startswith("• "):
            content = (inline_text or raw_text)[2:].strip()
            blocks.append(("list", f"{'  ' * list_level}- {content}"))
            continue

        ordered_match = _DOCX_ORDERED_MARKER_RE.match(raw_text)
        if ordered_match:
            marker = ordered_match.group("number")
            content_text = ordered_match.group("content")
            if inline_text:
                inline_match = _DOCX_ORDERED_MARKER_RE.match(inline_text)
                if inline_match:
                    marker = inline_match.group("number")
                    content_text = inline_match.group("content")
            blocks.append(("list", f"{'  ' * list_level}{marker}. {content_text.strip()}"))
            continue

        if _is_docx_blockquote(paragraph):
            quote_depth = max(1, round(paragraph.paragraph_format.left_indent.pt / 20))
            blocks.append(("blockquote", f"{'> ' * quote_depth}{(inline_text or raw_text).strip()}"))
            continue

        blocks.append(("paragraph", inline_text or raw_text))

    markdown_lines: list[str] = []
    previous_kind: str | None = None
    for kind, content in blocks:
        if markdown_lines:
            if previous_kind == "list" and kind == "list":
                markdown_lines.append(content)
            elif previous_kind == "blockquote" and kind == "blockquote":
                markdown_lines.append(content)
            else:
                markdown_lines.append("")
                markdown_lines.append(content)
        else:
            markdown_lines.append(content)
        previous_kind = kind

    return "\n".join(markdown_lines).strip()
